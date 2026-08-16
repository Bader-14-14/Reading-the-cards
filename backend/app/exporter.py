import os
from docx import Document
from openpyxl import Workbook
import zipfile
import tempfile
import os


def _display_fields(data: dict, language: str = 'ar') -> list[tuple[str, str]]:
    english = language.lower().startswith('en')
    number_key = 'iqama_number' if data.get('iqama_number') else 'id_number'
    labels = {
        'name': 'Name' if english else 'الاسم',
        number_key: 'Iqama Number' if english and number_key == 'iqama_number' else ('ID Number' if english else ('رقم الإقامة' if number_key == 'iqama_number' else 'رقم الهوية')),
        'nationality': 'Nationality' if english else 'الجنسية',
        'dob': 'Date of Birth' if english else 'تاريخ الميلاد',
        'doe': 'Date of Expiry' if english else 'تاريخ الانتهاء',
    }
    return [(labels[key], str(data[key])) for key in labels if data.get(key, '') not in ('', None)]


def create_word(data: dict, out_path: str, language: str = 'ar') -> str:
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc = Document()
    doc.add_heading('Document Reading Results' if language.lower().startswith('en') else 'نتائج قراءة الوثيقة', level=1)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field' if language.lower().startswith('en') else 'الحقل'
    hdr_cells[1].text = 'Value' if language.lower().startswith('en') else 'القيمة'
    for k, v in _display_fields(data, language):
        row_cells = table.add_row().cells
        row_cells[0].text = str(k)
        row_cells[1].text = str(v)
    doc.save(out_path)
    return out_path


def create_excel(data: dict, out_path: str, language: str = 'ar') -> str:
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.append(['Field', 'Value'] if language.lower().startswith('en') else ['الحقل', 'القيمة'])
    for k, v in _display_fields(data, language):
        ws.append([k, v])
    wb.save(out_path)
    return out_path


def create_excel_from_list(rows: list, out_path: str) -> str:
    """rows: list of dicts - each dict maps field->value. Create a table with columns as union of keys."""
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    wb = Workbook()
    ws = wb.active
    # collect all keys in order
    keys = []
    for r in rows:
        for k in r.keys():
            if k not in keys:
                keys.append(k)
    ws.append(keys)
    for r in rows:
        ws.append([r.get(k, '') for k in keys])
    wb.save(out_path)
    return out_path


def create_word_zip(rows: list, out_zip_path: str) -> str:
    """Create individual Word docs for each row and package into a zip."""
    os.makedirs(os.path.dirname(out_zip_path), exist_ok=True)
    with tempfile.TemporaryDirectory() as td:
        doc_paths = []
        for i, r in enumerate(rows, start=1):
            p = os.path.join(td, f'doc_{i}.docx')
            create_word(r, p)
            doc_paths.append(p)
        # create zip
        with zipfile.ZipFile(out_zip_path, 'w', zipfile.ZIP_DEFLATED) as z:
            for p in doc_paths:
                z.write(p, arcname=os.path.basename(p))
    return out_zip_path
